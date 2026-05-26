import os
import time
import subprocess
from pathlib import Path
from docx import Document
from pdfminer.high_level import extract_text
import platform


class TextExtractor:

    def __init__(self):
        pass

    # =========================
    # 获取文件后缀
    # =========================
    def _get_file_ext(self, file_path: str) -> str:
        return Path(file_path).suffix.lower()

    # =========================
    # 提取 DOC
    # Linux方案：
    # DOC -> DOCX -> 读取
    # =========================


    def _extract_doc(self, file_path: str) -> str:

        try:

            time.sleep(1)

            file_path = os.path.abspath(file_path)

            output_dir = os.path.dirname(file_path)

            # =========================
            # 判断系统
            # =========================
            system_name = platform.system()

            # Windows
            if system_name == "Windows":

                libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

                if not os.path.exists(libreoffice_path):
                    return "Windows 未安装 LibreOffice"

            # Linux / Ubuntu
            else:

                libreoffice_path = "libreoffice26.2"

            # =========================
            # DOC -> DOCX
            # =========================
            cmd = [
                libreoffice_path,
                "--headless",
                "--convert-to",
                "docx",
                file_path,
                "--outdir",
                output_dir
            ]

            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )

            if result.returncode != 0:
                return f"DOC 转换失败：{result.stderr}"

            # =========================
            # docx路径
            # =========================
            docx_path = str(Path(file_path).with_suffix(".docx"))

            if not os.path.exists(docx_path):
                return "DOC 转 DOCX 失败"

            # =========================
            # 读取 docx
            # =========================
            doc = Document(docx_path)

            text_list = []

            for para in doc.paragraphs:

                text = para.text.strip()

                if text:
                    text_list.append(text)

            return "\n".join(text_list)

        except Exception as e:
            return f"DOC 提取失败：{str(e)}"


    # def _extract_doc(self, file_path: str) -> str:

    #     try:

    #         time.sleep(1)

    #         file_path = os.path.abspath(file_path)

    #         output_dir = os.path.dirname(file_path)

    #                 # Windows 下 LibreOffice 路径
    #         libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

    #         # 判断是否存在
    #         if not os.path.exists(libreoffice_path):
    #             return "未找到 LibreOffice，请检查安装路径"

    #         # LibreOffice 转 DOCX
    #         cmd = [
    #             libreoffice_path,
    #             "--headless",
    #             "--convert-to",
    #             "docx",
    #             file_path,
    #             "--outdir",
    #             output_dir
    #         ]

    #         result = subprocess.run(
    #             cmd,
    #             stdout=subprocess.PIPE,
    #             stderr=subprocess.PIPE,
    #             text=True
    #         )

    #         if result.returncode != 0:
    #             return f"DOC 转换失败：{result.stderr}"

    #         # 生成 docx 路径
    #         docx_path = str(Path(file_path).with_suffix(".docx"))

    #         if not os.path.exists(docx_path):
    #             return "DOC 转 DOCX 失败"

    #         # 读取 DOCX
    #         doc = Document(docx_path)

    #         text_list = []

    #         for para in doc.paragraphs:

    #             text = para.text.strip()

    #             if text:
    #                 text_list.append(text)

    #         return "\n".join(text_list)

    #     except Exception as e:
    #         return f"DOC 提取失败：{str(e)}"

    # =========================
    # 提取 DOCX
    # =========================
    def _extract_docx(self, file_path: str) -> str:

        try:

            time.sleep(1)

            doc = Document(file_path)

            text_list = []

            for para in doc.paragraphs:

                text = para.text.strip()

                if text:
                    text_list.append(text)

            return "\n".join(text_list)

        except Exception as e:
            return f"DOCX 提取失败：{str(e)}"

    # =========================
    # 提取 PDF
    # =========================
    def _extract_pdf(self, file_path: str) -> str:

        try:

            time.sleep(1)

            text = extract_text(file_path)

            return text.strip()

        except Exception as e:
            return f"PDF 提取失败：{str(e)}"

    # =========================
    # 主入口
    # =========================
    def extract_text(self, file_path: str) -> str:

        if not os.path.exists(file_path):
            return "错误：文件不存在！"

        ext = self._get_file_ext(file_path)

        # DOC
        if ext == ".doc":
            return self._extract_doc(file_path)

        # DOCX
        elif ext == ".docx":
            return self._extract_docx(file_path)

        # PDF
        elif ext == ".pdf":
            return self._extract_pdf(file_path)

        else:
            return f"不支持格式：{ext}"


# =========================
# 使用示例
# =========================
if __name__ == "__main__":

    extractor = TextExtractor()

    test_files = [
        "/home/test/1.doc",
        "/home/test/2.docx",
        "/home/test/3.pdf"
    ]

    for file_path in test_files:

        print(f"\n📄 提取文件: {file_path}")

        print("-" * 60)

        result = extractor.extract_text(file_path)

        print(result)