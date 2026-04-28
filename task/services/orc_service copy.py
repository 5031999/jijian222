import os
from pathlib import Path
from PIL import Image
import pytesseract
from docx import Document
from pdfminer.high_level import extract_text
import textract
import time

# ===================== 关键修复：正确配置 antiword 路径 =====================
# 把 antiword.exe 所在文件夹加入系统 PATH（最稳定、不会报错的方法）
antiword_folder = r"D:\\antiword-0.37\\antiword-0_37-windows\\antiword"
os.environ["PATH"] = antiword_folder + os.pathsep + os.environ["PATH"]

# 同时设置 HOME（antiword 必须）
os.environ["HOME"] = antiword_folder


class TextExtractor:

    def __init__(self, tesseract_path=None):
        if tesseract_path:
            if not os.path.exists(tesseract_path):
                raise ValueError(f"Tesseract 路径不存在：{tesseract_path}")
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

    def _get_file_ext(self, file_path: str) -> str:
        return Path(file_path).suffix.lower()

    def _extract_doc(self, file_path: str) -> str:
        """使用 textract 提取老式 DOC 文本（已修复路径问题）"""
        try:
            time.sleep(1)
            #新版 textract 直接调用即可，无需修改内部属性
            text = textract.process(file_path, encoding='utf-8')
            return text.decode("utf-8", errors="ignore")
            #return f"doc文件"
        except Exception as e:
            return f"doc提取失败：{str(e)}"

    def _extract_docx(self, file_path: str) -> str:
        """使用 python-docx 提取 DOCX 文本"""
        try:
            time.sleep(1)
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"docx提取失败：{str(e)}"

    def _extract_pdf(self, file_path: str) -> str:
        """使用 pdfminer.six 提取 PDF 文本"""
        try:
            time.sleep(1)
            text = extract_text(file_path)
            return text.strip()
        except Exception as e:
            return f"PDF提取失败：{str(e)}"

    def _extract_image(self, file_path: str) -> str:
        """使用 pytesseract 提取图片文字"""
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            return text.strip()
        except Exception as e:
            return f"图片OCR失败：{str(e)}"

    def extract_text(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return "错误：文件不存在！"
        ext = self._get_file_ext(file_path)
        if ext == ".doc":
            return self._extract_doc(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        elif ext == ".pdf":
            return self._extract_pdf(file_path)
        # elif ext in [".jpg", ".jpeg", ".png", ".bmp"]:
        #     return self._extract_image(file_path)
        else:
            return f"不支持格式：{ext}"


# ===================== 使用示例 =====================
if __name__ == "__main__":
    extractor = TextExtractor(
        tesseract_path=r"D:\soft\ocr\tesseract.exe"
    )

    # 测试 DOC
    for file_path in [
        r"D:\\1.doc"
    ]:
        print(f"\n📄 提取文件: {file_path}")
        print("-" * 60)
        print(extractor.extract_text(file_path))

# import os
# from pathlib import Path
# from PIL import Image
# import pytesseract
# from docx import Document  # ✅ 处理 docx
# import pdfplumber           # ✅ 处理 PDF

# class TextExtractor:

#     def __init__(self, tesseract_path=None):
#         """
#         tesseract_path: tesseract.exe 路径
#         """
#         if tesseract_path:
#             if not os.path.exists(tesseract_path):
#                 raise ValueError(f"Tesseract 路径不存在：{tesseract_path}")
#             pytesseract.pytesseract.tesseract_cmd = tesseract_path

#     def _get_file_ext(self, file_path: str) -> str:
#         return Path(file_path).suffix.lower()

#     def _extract_docx(self, file_path: str) -> str:
#         """
#         使用 python-docx 提取 docx 文本
#         """
#         try:
#             doc = Document(file_path)
#             return "\n".join([para.text for para in doc.paragraphs])
#         except Exception as e:
#             return f"docx提取失败：{str(e)}"

#     def _extract_pdf(self, file_path: str) -> str:
#         """
#         使用 pdfplumber 提取 PDF 文本
#         """
#         try:
#             text = ""
#             with pdfplumber.open(file_path) as pdf:
#                 for page in pdf.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text + "\n"
#             return text.strip()
#         except Exception as e:
#             return f"PDF提取失败：{str(e)}"

#     def _extract_image(self, file_path: str) -> str:
#         """
#         使用 pytesseract 提取图片文字
#         """
#         try:
#             img = Image.open(file_path)
#             text = pytesseract.image_to_string(img, lang="chi_sim+eng")
#             return text.strip()
#         except Exception as e:
#             return f"图片OCR失败：{str(e)}"

#     def extract_text(self, file_path: str) -> str:
#         if not os.path.exists(file_path):
#             return "错误：文件不存在！"
#         ext = self._get_file_ext(file_path)
#         if ext == ".docx":
#             return self._extract_docx(file_path)
#         elif ext == ".pdf":
#             return self._extract_pdf(file_path)
#         elif ext in [".jpg", ".jpeg", ".png", ".bmp"]:
#             return self._extract_image(file_path)
#         else:
#             return f"不支持格式：{ext}"


# # ===================== 使用示例 =====================
# if __name__ == "__main__":
#     extractor = TextExtractor(
#         tesseract_path=r"D:\\soft\\ocr\\tesseract.exe"
#     )

#     # 测试 DOCX / PDF / 图片
#     for file_path in [
#         r"D:\\1.pdf"
#     ]:
#         print(f"\n📄 提取文件: {file_path}")
#         print("-" * 60)
#         print(extractor.extract_text(file_path))